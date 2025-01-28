from fastapi import FastAPI, HTTPException, Request, Form
from fastapi.responses import HTMLResponse, FileResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from docx import Document
import os
import datetime
import subprocess
from collections import Counter
import re


app = FastAPI()

# Configuração de templates e arquivos estáticos
templates = Jinja2Templates(directory="templates")
app.mount("/static", StaticFiles(directory="static"), name="static")

# Diretório para salvar os resumos
OUTPUT_DIR = "resumos"
os.makedirs(OUTPUT_DIR, exist_ok=True)

@app.get("/", response_class=HTMLResponse)
async def home(request: Request):
    """
    Página inicial da aplicação.
    """
    return templates.TemplateResponse("index.html", {"request": request})


@app.post("/summarize/")
async def summarize_text(request: Request, text: str = Form(...)):
    """
    Endpoint para resumir o texto e salvar no Word.
    """
    try:
        # Resumir o texto usando o modelo LLaMA
        summary = run_llama(text)

        # Salvar o resumo em um arquivo Word
        file_path = save_to_word(summary)

        return templates.TemplateResponse(
            "index.html",
            {"request": request, "summary": summary, "file_path": file_path},
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/download/{file_name}")
async def download_file(file_name: str):
    """
    Permite o download de um arquivo Word gerado.
    """
    file_path = os.path.join(OUTPUT_DIR, file_name)
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="Arquivo não encontrado.")
    return FileResponse(file_path, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


def run_llama(text):
    """
    Executa o modelo LLaMA via Ollama para resumir o texto.
    """
    try:
        # Comando para rodar o modelo usando Ollama
        command = ["ollama", "run", "llama3.2:1b", f"Summarize: {text}"]
        result = subprocess.run(command, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True, encoding="utf-8")

        if result.returncode != 0:
            raise Exception(f"Ollama error: {result.stderr.strip()}")

        # Log da saída do comando
        print(f"Saída do comando: {result.stdout.strip()}")

        # Retornar a saída diretamente (ou ajustar conforme necessário)
        return result.stdout.strip()
    except Exception as e:
        raise Exception(f"Erro ao executar o modelo: {str(e)}")
    

def extract_keyword(summary):
    """
    Extrai a palavra-chave mais relevante do resumo.
    """
    # Remover caracteres especiais e transformar em minúsculas
    words = re.findall(r'\b\w+\b', summary.lower())
    
    # Ignorar palavras comuns (stopwords)
    stopwords = {"de", "da", "do", "em", "e", "o", "a", "os", "as", "que", "é", "para", "com", "um", "uma", "por"}
    filtered_words = [word for word in words if word not in stopwords and len(word) > 2]

    # Encontrar a palavra mais frequente
    most_common = Counter(filtered_words).most_common(1)
    return most_common[0][0] if most_common else "resumo"

def generate_friendly_title(summary):
    """
    Gera um título amigável baseado no resumo.
    """
    # Extrair palavra-chave
    keyword = extract_keyword(summary)
    
    # Criar um título amigável
    friendly_title = f"Resumo sobre {keyword.capitalize()}"
    return friendly_title

def save_to_word(summary):
    """
    Salva o resumo em um arquivo Word com um nome amigável.
    """
    # Garantir que o diretório OUTPUT_DIR exista
    if not os.path.exists(OUTPUT_DIR):
        os.makedirs(OUTPUT_DIR)

    # Criar o documento Word
    doc = Document()
    doc.add_heading("Resumo Gerado", level=1)
    doc.add_paragraph(summary)

    # Gerar título amigável
    friendly_title = generate_friendly_title(summary)

    # Nome do arquivo com o título amigável
    file_name = f"{friendly_title.replace(' ', '_')}_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    file_path = os.path.join(OUTPUT_DIR, file_name)

    # Salvar o documento
    doc.save(file_path)
    return file_name

